"""文档解析：PDF / DOCX / XLSX / TXT / MD -> (页码|段落, 文本) 单元列表。

- PDF 保留页码；扫描件/无文本层直接报错（本试点不含 OCR）。
- DOCX/TXT/MD 以“段落号”作为位置引用；XLSX 以工作表行作为文本单元。
"""
from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".txt": "txt",
    ".md": "md",
}

# 各类型的允许 MIME；application/octet-stream 一律放行，最终以魔数+实际解析为准
MIME_MAP: dict[str, set[str]] = {
    "pdf": {"application/pdf", "application/x-pdf"},
    "docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "xlsx": {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
    "txt": {"text/plain"},
    "md": {"text/markdown", "text/x-markdown", "text/plain"},
}
_ALLOWED_MIMES: set[str] = set().union(*MIME_MAP.values()) | {"application/octet-stream"}


class ParseError(Exception):
    def __init__(self, message: str, code: str = "parse_error") -> None:
        super().__init__(message)
        self.message = message
        self.code = code


@dataclass
class Unit:
    """一段可独立切块的文本；page/paragraph 至少一个用于来源引用。"""

    text: str
    page: int | None = None          # 1 起始页码（仅 PDF）
    paragraph: int | None = None     # 1 起始段落号（DOCX/TXT/MD）


def detect_ext(filename: str) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext not in EXTENSIONS:
        raise ParseError(
            f"不支持的文件类型 {ext or '(无扩展名)'}：仅支持 PDF / DOCX / XLSX / TXT / MD",
            code="unsupported_ext",
        )
    return ext


def check_mime(content_type: str) -> None:
    ctype = (content_type or "").split(";")[0].strip().lower()
    if ctype and ctype not in _ALLOWED_MIMES:
        raise ParseError(f"MIME 类型 {ctype!r} 不在允许范围，疑似伪造扩展名", code="bad_mime")


def check_magic(kind: str, data: bytes) -> None:
    if kind == "pdf":
        if not data[:5].startswith(b"%PDF-"):
            raise ParseError("文件内容不是有效 PDF（缺少 %PDF- 文件头）", code="bad_magic")
    elif kind == "docx":
        if data[:4] != b"PK\x03\x04":
            raise ParseError("文件内容不是有效 DOCX（缺少 ZIP/OOXML 头）", code="bad_magic")
        try:
            names = zipfile.ZipFile(io.BytesIO(data)).namelist()
        except zipfile.BadZipFile as exc:
            raise ParseError(f"文件内容不是有效 DOCX（ZIP 损坏）：{exc}", code="bad_magic") from exc
        if "word/document.xml" not in names:
            raise ParseError("文件内容不是有效 DOCX（缺少 word/document.xml）", code="bad_magic")
    elif kind == "xlsx":
        if data[:4] != b"PK\x03\x04":
            raise ParseError("文件内容不是有效 XLSX（缺少 ZIP/OOXML 头）", code="bad_magic")
        try:
            names = zipfile.ZipFile(io.BytesIO(data)).namelist()
        except zipfile.BadZipFile as exc:
            raise ParseError(f"文件内容不是有效 XLSX（ZIP 损坏）：{exc}", code="bad_magic") from exc
        if "xl/workbook.xml" not in names:
            raise ParseError("文件内容不是有效 XLSX（缺少 xl/workbook.xml）", code="bad_magic")


def _clean(text: str) -> str:
    text = (
        text.replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\x0c", "\n")
        .replace("\u3000", " ")
    )
    out: list[str] = []
    blank = False
    for line in text.split("\n"):
        line = line.rstrip()
        if not line.strip():
            if out and not blank:
                out.append("")
            blank = True
        else:
            out.append(line)
            blank = False
    text = "\n".join(out)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


# ---------------- PDF ----------------

def parse_pdf(path: Path) -> list[Unit]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise ParseError("PDF 已加密且无法解密", code="pdf_encrypted") from exc
    units: list[Unit] = []
    total = 0
    for idx, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            raise ParseError(f"第 {idx} 页文本提取失败：{exc}", code="pdf_extract") from exc
        text = _clean(text)
        total += len(text)
        if text:
            units.append(Unit(text=text, page=idx))
    if total == 0:
        raise ParseError(
            "PDF 中未提取到任何文本（可能是扫描件/图片型 PDF）。本试点不含 OCR，"
            "请改用带文字层的 PDF。",
            code="scanned_pdf",
        )
    return units


# ---------------- DOCX ----------------

def parse_docx(path: Path) -> list[Unit]:
    from docx import Document

    doc = Document(str(path))
    units: list[Unit] = []
    para_no = 0

    def add(text: str) -> None:
        nonlocal para_no
        text = _clean(text)
        if text:
            para_no += 1
            units.append(Unit(text=text, paragraph=para_no))

    for p in doc.paragraphs:
        add(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                add(cell.text)
    if not units:
        raise ParseError("DOCX 中没有可索引的文本内容", code="empty_doc")
    return units


# ---------------- XLSX ----------------

def parse_xlsx(path: Path) -> list[Unit]:
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter

    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        raise ParseError(f"XLSX 读取失败：{exc}", code="xlsx_extract") from exc

    units: list[Unit] = []
    try:
        for sheet in workbook.worksheets:
            for row_no, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells = []
                for col_no, value in enumerate(row, start=1):
                    if value is None:
                        continue
                    text = _clean(str(value))
                    if text:
                        cells.append(f"{get_column_letter(col_no)}{row_no}={text}")
                if cells:
                    units.append(
                        Unit(
                            text=f"工作表《{sheet.title}》第 {row_no} 行：" + "；".join(cells),
                            paragraph=row_no,
                        )
                    )
    finally:
        workbook.close()
    if not units:
        raise ParseError("XLSX 中没有可索引的单元格内容", code="empty_doc")
    return units


# ---------------- TXT / MD ----------------

def _read_text_bytes(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ParseError("文本编码无法识别（支持 UTF-8 / GB18030）", code="bad_encoding")


def parse_text(path: Path) -> list[Unit]:
    text = _clean(_read_text_bytes(path))
    units: list[Unit] = []
    para_no = 0
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        para_no += 1
        units.append(Unit(text=para, paragraph=para_no))
    if not units:
        raise ParseError("文件内容为空，无可索引文本", code="empty_doc")
    return units


PARSERS: dict[str, object] = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "xlsx": parse_xlsx,
    "txt": parse_text,
    "md": parse_text,
}
