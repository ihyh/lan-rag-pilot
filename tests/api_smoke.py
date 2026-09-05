"""端到端冒烟测试（需应用与 mock DeepSeek 均已启动，见 docs 测试手册）。

先准备两个进程：
  1) mock 模型:      uvicorn tests.mock_deepseek:app --port 8099
  2) 应用(离线模式):  按 tests/smoke_runner.sh 内的环境变量设置后
     uvicorn app.main:app --port 8090
     （Linux/Docker 环境下可直接：docker run --rm -v "$PWD:/rag" -w /rag <含依赖镜像> bash tests/smoke_runner.sh）

然后：.venv\\Scripts\\python tests\\api_smoke.py
覆盖测试计划中的：登录/错误密码/注销、user 权限 403、root 全部管理、
四种格式（使用脚本生成的最小有效样本）、扫描 PDF、重复文件、超限、伪造扩展名、空文件、
问答与来源、历史可见性、限流 429、LLM 异常不泄露(HTTP 层)、审计。
"""
from __future__ import annotations

import io
import os
import sqlite3
import sys
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import httpx

BASE_URL = os.environ.get("RAG_SMOKE_URL", "http://127.0.0.1:8090")
ROOT_PW = os.environ.get("RAG_ROOT_PASSWORD", "fcd123")

PASS: list[str] = []
FAIL: list[str] = []


def check(cond: bool, msg: str) -> None:
    if cond:
        PASS.append(msg)
        print(f"  [PASS] {msg}")
    else:
        FAIL.append(msg)
        print(f"  [FAIL] {msg}")


def err_text(resp: httpx.Response) -> str:
    try:
        data = resp.json()
        detail = data.get("detail")
        if isinstance(detail, dict):
            return f"{detail.get('code','')}:{detail.get('message','')}"
        return str(detail)
    except Exception:
        return resp.text[:200]


class Smoke:
    def __init__(self) -> None:
        self.c = httpx.Client(base_url=BASE_URL, timeout=60.0)
        self.root: dict = {}
        self.alice: dict = {}
        self.bob: dict = {}

    def login(self, username: str, password: str) -> httpx.Response:
        return self.c.post("/api/login", json={"username": username, "password": password})

    # ---------- 1. 认证 ----------
    def test_auth(self) -> None:
        print("\n== 认证 ==")
        r = self.c.get("/login")
        check(r.status_code == 200 and "登录" in r.text, "登录页面可访问")
        r = self.c.get("/static/manifest.webmanifest")
        check(r.status_code == 200 and r.json().get("name"), "Web App Manifest 可访问")
        r = self.c.get("/static/icons/icon.svg")
        check(r.status_code == 200 and "<svg" in r.text, "应用图标可访问")
        r = self.c.get("/sw.js")
        check(
            r.status_code == 200 and r.headers.get("service-worker-allowed") == "/",
            "Service Worker 入口及作用域正确",
        )
        r = self.login("root", "wrong-password")
        check(r.status_code == 401, "错误密码返回 401")
        r = self.login("root", ROOT_PW)
        check(r.status_code == 200, "root/fcd123 可登录")
        self.root = r.json()["user"]
        check(self.root.get("role") == "root", "root 角色正确")
        r = self.c.get("/api/me")
        check(r.status_code == 200 and r.json().get("username") == "root", "/api/me 正常")
        # 密码不以明文落库：登录用任意大密码不应成功且不泄露
        r = self.c.get("/api/health")
        check(r.status_code == 200, "健康检查 200")
        r = self.c.get("/api/ready")
        check(r.status_code == 200 and r.json().get("status") == "ready", "就绪检查 200")
        with sqlite3.connect(os.environ["RAG_DB_PATH"]) as db:
            password_hash = db.execute(
                "SELECT password_hash FROM users WHERE username='root'"
            ).fetchone()[0]
            session_hash = db.execute(
                "SELECT token_hash FROM sessions ORDER BY id DESC LIMIT 1"
            ).fetchone()[0]
        check(password_hash.startswith("$argon2") and password_hash != ROOT_PW, "root 密码仅以 Argon2 哈希保存")
        check(session_hash != self.c.cookies.get("rag_session"), "会话令牌仅以哈希保存")

    # ---------- 2. user 权限 ----------
    def test_user_permissions(self) -> None:
        print("\n== user 权限（403）==")
        r = self.c.post(
            "/api/admin/users",
            json={"username": "alice", "password": "alice123", "role": "user"},
        )
        check(r.status_code == 201, "root 创建 user 成功")
        r = self.login("alice", "alice123")
        check(r.status_code == 200, "alice 可登录")
        self.alice = r.json()["user"]

        r = self.c.post("/api/admin/documents", files={"file": ("x.txt", b"data", "text/plain")})
        check(r.status_code == 403, "user 上传返回 403")
        r = self.c.post("/api/admin/users", json={"username": "x", "password": "x123456", "role": "user"})
        check(r.status_code == 403, "user 建账号返回 403")
        r = self.c.get("/api/admin/audit")
        check(r.status_code == 403, "user 看审计返回 403")
        r = self.c.get("/api/admin/feedback.csv")
        check(r.status_code == 403, "user 导出反馈返回 403")
        r = self.c.post("/api/admin/documents/1/reindex")
        check(r.status_code == 403, "user 重建索引返回 403")
        r = self.c.delete("/api/admin/documents/1")
        check(r.status_code == 403, "user 删除文档返回 403")
        r = self.c.get("/api/admin/users")
        check(r.status_code == 403, "user 管理用户返回 403")
        r = self.c.get("/api/admin/documents")
        check(r.status_code == 403, "user 查看管理文档返回 403")
        for path in ["/api/admin/departments", "/api/admin/knowledge-bases", "/api/knowledge-bases"]:
            check(self.c.get(path).status_code == 404, f"分类接口已移除：{path}")

    # ---------- 3. 文档入库（格式/重复/伪造/空/超限） ----------
    @staticmethod
    def _txt_bytes() -> bytes:
        paras = []
        for i in range(24):
            paras.append(
                f"第 {i + 1} 节 局域网考勤管理规定：员工应在工作日 9:00 前完成打卡登记，"
                "迟到需在系统内填写迟到说明并由直属主管审批；出差前需提交申请，"
                "市内交通凭票据据实报销，出差住宿每日上限 400 元。"
            )
        return ("\n\n".join(paras)).encode("utf-8")

    @staticmethod
    def _docx_bytes() -> bytes:
        from docx import Document

        doc = Document()
        doc.add_paragraph("局域网运维手册（测试文档）")
        for i in range(30):
            doc.add_paragraph(
                f"第 {i + 1} 条：机房温度应保持在 18-27 摄氏度；UPS 电池每季度巡检一次；"
                "核心交换机配置变更需提前一天审批并留存审计记录。"
            )
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _xlsx_bytes() -> bytes:
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "设备清单"
        sheet.append(["设备", "巡检周期", "负责人"])
        sheet.append(["UPS", "每季度", "运维组"])
        sheet.append(["核心交换机", "每月", "网络组"])
        buf = io.BytesIO()
        workbook.save(buf)
        workbook.close()
        return buf.getvalue()

    @staticmethod
    def _pdf_bytes() -> bytes:
        text = "Pilot policy: hotel reimbursement limit is 400 yuan per day."
        stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
            b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        ]
        out = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, obj in enumerate(objects, start=1):
            offsets.append(len(out))
            out.extend(f"{index} 0 obj\n".encode("ascii") + obj + b"\nendobj\n")
        xref = len(out)
        out.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
        for offset in offsets[1:]:
            out.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        out.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii"))
        return bytes(out)

    @staticmethod
    def _blank_pdf_bytes() -> bytes:
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue()

    def test_documents(self) -> None:
        print("\n== 文档管理（root）==")
        r = self.login("root", ROOT_PW)
        check(r.status_code == 200, "文档管理前重新登录 root")
        txt = self._txt_bytes()
        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("考勤与报销规定.txt", txt, "text/plain")},
            data={"version": "2026.1", "effective_date": "2026-09-01", "tags": "研发,制度,研发"},
        )
        check(r.status_code == 201 and r.json().get("status") == "ready", "TXT 上传并入库成功")
        self.txt_doc = r.json()
        check(int(self.txt_doc["num_chunks"]) >= 1, f"TXT 切片数 {self.txt_doc['num_chunks']} >= 1")
        check(
            self.txt_doc.get("version") == "2026.1"
            and self.txt_doc.get("effective_date") is None
            and self.txt_doc.get("tags") == [],
            "文档版本写入，客户端提交的生效日期和标签被忽略",
        )
        upload_date = self.txt_doc["created_at"][:10]
        r = self.c.get("/api/admin/documents", params={"version": "2026.1", "uploaded_date_from": upload_date, "uploaded_date_to": upload_date})
        check(r.status_code == 200 and [d["id"] for d in r.json()["items"]] == [self.txt_doc["id"]], "文档按版本与上传日期筛选成功")
        r = self.c.get("/api/admin/documents", params={"uploaded_date_from": "2026-10-01", "uploaded_date_to": "2026-09-01"})
        check(r.status_code == 422, "非法上传日期范围返回 422")
        r = self.c.get(f"/api/documents/{self.txt_doc['id']}/file")
        check(
            r.status_code == 200 and r.content == txt and "inline" in r.headers.get("content-disposition", ""),
            "root 可内联打开原文",
        )
        docx = self._docx_bytes()
        r = self.c.post(
            "/api/admin/documents",
            files={
                "file": (
                    "运维手册.docx",
                    docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        check(r.status_code == 201 and r.json().get("status") == "ready", "DOCX 上传并入库成功")
        self.docx_doc = r.json()

        r = self.c.post(
            "/api/admin/documents",
            files={
                "file": (
                    "设备清单.xlsx",
                    self._xlsx_bytes(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )
        check(r.status_code == 201 and r.json().get("status") == "ready", "XLSX 上传并入库成功")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("制度说明.md", b"# Knowledge base\n\nMarkdown pilot document with unique content.", "text/markdown")},
        )
        check(r.status_code == 201 and r.json().get("status") == "ready", "Markdown 上传并入库成功")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("报销规则.pdf", self._pdf_bytes(), "application/pdf")},
        )
        check(r.status_code == 201 and r.json().get("pages") == 1, "文本 PDF 上传并保留页码")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("扫描件.pdf", self._blank_pdf_bytes(), "application/pdf")},
        )
        check(r.status_code == 400 and "OCR" in err_text(r), "无文本 PDF 明确提示不含 OCR")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("考勤与报销规定-copy.txt", txt, "text/plain")},
        )
        check(r.status_code == 409, "重复文件(SHA-256)返回 409")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("伪装.pdf", txt, "application/pdf")},
        )
        check(r.status_code == 400 and "PDF" in err_text(r), "伪造扩展名返回 400")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("伪装.xlsx", txt, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        check(r.status_code == 400 and "XLSX" in err_text(r), "伪造 XLSX 扩展名返回 400")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("错误类型.txt", b"plain text", "image/png")},
        )
        check(r.status_code == 400 and "MIME" in err_text(r), "错误 MIME 返回 400")

        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("空文件.txt", b"", "text/plain")},
        )
        check(r.status_code == 400, "空文件返回 400")

        big = b"x" * (25 * 1024 * 1024 + 1024)
        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("超大文件.txt", big, "text/plain")},
        )
        check(r.status_code == 413, "超过 25MB 返回 413")

        r = self.c.get("/api/admin/documents")
        check(r.status_code == 200 and len(r.json()["items"]) >= 2, "文档列表包含已入库文件")

        # 文档删除后再查列表
        target = self.docx_doc["id"]
        r = self.c.delete(f"/api/admin/documents/{target}")
        check(r.status_code == 204, "root 删除文档 204")
        r = self.c.get("/api/admin/documents")
        check(all(d["id"] != target for d in r.json()["items"]), "删除后列表同步")

        # 重新索引（重新上传一个再测）
        r = self.c.post(
            "/api/admin/documents",
            files={"file": ("运维手册2.docx", self._docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc2 = r.json()
        r = self.c.post(f"/api/admin/documents/{doc2['id']}/reindex")
        check(r.status_code == 200 and r.json().get("status") == "ready", "重新索引成功")

    def test_kb_admin_permissions(self) -> None:
        print("\n== 统一文档库与文档管理员权限 ==")
        keeper_response = self.c.post(
            "/api/admin/users",
            json={"username": "keeper", "password": "keeper123", "role": "kb_admin"},
        )
        check(keeper_response.status_code == 201, "root 可创建文档管理员")
        keeper = keeper_response.json()
        # 模拟旧数据库的跨部门数据：保留旧归属但不再据此限制访问。
        with sqlite3.connect(os.environ["RAG_DB_PATH"]) as db:
            stored = db.execute("SELECT role, is_kb_admin FROM users WHERE id=?", (keeper["id"],)).fetchone()
            for dep_id in (10, 11):
                db.execute("INSERT INTO departments (id,name,created_at,updated_at) VALUES (?,?,?,?)",
                           (dep_id, f"旧部门{dep_id}", "2026-01-01", "2026-01-01"))
                db.execute("INSERT INTO knowledge_bases (id,name,department_id,created_at,updated_at) VALUES (?,?,?,?,?)",
                           (dep_id, f"旧分类{dep_id}", dep_id, "2026-01-01", "2026-01-01"))
            db.execute("INSERT INTO user_departments VALUES (?,?)", (keeper["id"], 10))
            db.execute("INSERT INTO document_knowledge_bases VALUES (?,?)", (self.txt_doc["id"], 11))
        check(stored == ("user", 1), "文档管理员兼容旧数据库角色")

        with httpx.Client(base_url=BASE_URL, timeout=60.0) as client:
            r = client.post("/api/login", json={"username": "keeper", "password": "keeper123"})
            check(r.status_code == 200 and r.json()["user"]["role"] == "kb_admin", "文档管理员可登录")
            check(client.get("/admin").status_code == 200, "文档管理员可进入后台")
            for path in ["/api/admin/users", "/api/admin/audit", "/api/admin/overview", "/api/admin/settings", "/api/admin/chats"]:
                check(client.get(path).status_code == 403, f"文档管理员仍不可访问 {path}")
            expected_ids = {d["id"] for d in self.c.get("/api/admin/documents").json()["items"]}
            actual_ids = {d["id"] for d in client.get("/api/admin/documents").json()["items"]}
            check(actual_ids == expected_ids, "文档管理员可查看全部文档，不受旧部门归属限制")
            r = client.get(f"/api/documents/{self.txt_doc['id']}/file")
            check(r.status_code == 200 and r.content == self._txt_bytes(), "文档管理员可打开原属其它部门的文档")
            r = client.post(f"/api/admin/documents/{self.txt_doc['id']}/reindex")
            check(r.status_code == 200, "文档管理员可重新处理原属其它部门文档")
            r = client.post("/api/admin/documents",
                            files={"file": ("统一文档.txt", "统一文档库无需选择部门。".encode("utf-8"), "text/plain")})
            check(r.status_code == 201, "文档管理员可直接上传文档")
            doc_id = r.json()["id"]
            with sqlite3.connect(os.environ["RAG_DB_PATH"]) as db:
                db.execute("INSERT INTO document_knowledge_bases VALUES (?,?)", (doc_id, 11))
            r = client.delete(f"/api/admin/documents/{doc_id}")
            check(r.status_code == 204, "文档管理员可删除原属其它部门文档")
            check(client.get(f"/api/documents/{doc_id}/file").status_code == 404, "已删除文档不可访问")
            check(client.post("/api/admin/users", json={"username":"forbidden","password":"pass123","role":"root"}).status_code == 403,
                  "文档管理员不能创建系统管理员")
        r = self.c.patch(f"/api/admin/users/{keeper['id']}", json={"department_ids": [11]})
        check(r.status_code == 422, "已移除的部门分配字段不会被静默接受")
        for path in ["/api/admin/departments", "/api/admin/knowledge-bases"]:
            check(self.c.post(path, json={"name":"已停用"}).status_code == 404, "旧分类创建接口已停用")
        check(self.c.patch(f"/api/admin/documents/{self.txt_doc['id']}/knowledge-bases",
                           json={"knowledge_base_ids":[10]}).status_code == 404, "旧文档范围分配接口已停用")

    # ---------- 4. 问答 ----------
    def test_query(self) -> None:
        print("\n== 问答 ==")
        r = self.c.post("/api/query", json={"question": "出差住宿上限是多少？"})
        check(r.status_code == 200, "root 问答成功")
        body = r.json()
        check(body.get("answer") and "Mock 模型" in body["answer"], "返回答案")
        check(len(body.get("sources", [])) > 0, f"来源数 {len(body.get('sources', []))} > 0")
        src = body["sources"][0]
        check(src.get("filename") and ("page" in src or "paragraph" in src), "来源含文件名与位置")
        self.root_chat_id = body["chat_id"]

        # 普通用户没有部门记录，仍可检索原属其它部门的文档。
        r = self.login("alice", "alice123")
        with sqlite3.connect(os.environ["RAG_DB_PATH"]) as db:
            n = db.execute("SELECT COUNT(*) FROM user_departments ud JOIN users u ON u.id=ud.user_id WHERE u.username='alice'").fetchone()[0]
        check(n == 0, "新用户无需部门分配")
        r = self.c.post("/api/query", json={"question": "出差住宿上限是多少？"})
        check(r.status_code == 200 and r.json().get("sources"), "无部门用户可检索统一文档库")
        r = self.c.get(f"/api/documents/{self.txt_doc['id']}/file")
        check(r.status_code == 200 and r.content == self._txt_bytes(), "普通用户可打开原属其它部门的原文")

        # user 也可问答
        r = self.login("alice", "alice123")
        check(r.status_code == 200, "问答前切换为 alice(user)")
        r = self.c.get(f"/api/documents/{self.txt_doc['id']}/file")
        check(r.status_code == 200 and r.content == self._txt_bytes(), "user 可打开统一文档库原文")
        r = self.c.post("/api/query", json={"question": "打卡时间是几点？"})
        check(r.status_code == 200, "alice(user) 问答成功")
        self.alice_chat_id = r.json()["chat_id"]
        r = self.c.post(f"/api/chats/{self.alice_chat_id}/feedback", json={"rating": "helpful"})
        check(r.status_code == 200 and r.json().get("rating") == "helpful", "user 提交回答反馈成功")
        r = self.c.get(f"/api/chats/{self.alice_chat_id}")
        check(r.status_code == 200 and r.json().get("feedback", {}).get("rating") == "helpful", "问答详情包含本人反馈")
        r = self.c.post(f"/api/chats/{self.root_chat_id}/feedback", json={"rating": "unhelpful"})
        check(r.status_code == 404, "user 不能评价他人问答")

    def test_history(self) -> None:
        print("\n== 历史与可见性 ==")
        r = self.c.get("/api/chats")
        items = r.json()["items"]
        check(r.status_code == 200 and any(c["id"] == self.alice_chat_id for c in items), "本人历史可见")
        r = self.c.get(f"/api/chats/{self.root_chat_id}")
        check(r.status_code == 404, "user 看不到他人问答记录(404)")
        r = self.c.delete(f"/api/chats/{self.root_chat_id}")
        check(r.status_code == 404, "user 不能删除他人问答记录(404)")
        r = self.c.get(f"/api/chats/{self.alice_chat_id}")
        check(r.status_code == 200 and len(r.json().get("sources", [])) > 0, "本人问答详情含来源")
        r = self.c.post("/api/query", json={"question": "这是一条用于删除测试的问答"})
        deleted_id = r.json().get("chat_id") if r.status_code == 200 else None
        check(r.status_code == 200 and deleted_id, "user 创建待删除问答")
        with sqlite3.connect(os.environ["RAG_DB_PATH"]) as db:
            source_count = db.execute(
                "SELECT COUNT(id) FROM chat_sources WHERE chat_id=?", (deleted_id,)
            ).fetchone()[0]
        check(source_count > 0, "待删除问答含关联引用")
        r = self.c.delete(f"/api/chats/{deleted_id}")
        check(r.status_code == 204, "user 可删除本人问答")
        r = self.c.get(f"/api/chats/{deleted_id}")
        check(r.status_code == 404, "已删除问答不可再读取")
        with sqlite3.connect(os.environ["RAG_DB_PATH"]) as db:
            source_count = db.execute(
                "SELECT COUNT(id) FROM chat_sources WHERE chat_id=?", (deleted_id,)
            ).fetchone()[0]
        check(source_count == 0, "删除问答同时删除关联引用")
        # 登出 alice 后会话失效
        r = self.c.post("/api/logout")
        check(r.status_code == 200, "注销成功")
        r = self.c.get("/api/me")
        check(r.status_code == 401, "注销后会话立即失效")
        r = self.c.get(f"/api/documents/{self.txt_doc['id']}/file")
        check(r.status_code == 401, "未登录不能打开原文")

    # ---------- 5. root 审计与全局可见 ----------
    def test_admin_views(self) -> None:
        print("\n== root 审计 ==")
        r = self.login("root", ROOT_PW)
        check(r.status_code == 200, "重新登录 root")
        r = self.c.get("/api/admin/chats")
        check(r.status_code == 200 and any(c["id"] == self.root_chat_id for c in r.json()["items"]), "root 可见全部问答")
        r = self.c.get("/api/admin/audit")
        acts = [a["action"] for a in r.json()["items"]]
        check("llm_query" in acts and "doc_upload" in acts and "document_open" in acts and "feedback_submit" in acts and "login" in acts, "审计含 llm_query/doc_upload/document_open/feedback/login")
        r = self.c.get("/api/admin/feedback")
        check(r.status_code == 200 and any(f["chat_id"] == self.alice_chat_id for f in r.json()["items"]), "root 可查看用户反馈")
        r = self.c.get("/api/admin/feedback.csv")
        check(r.status_code == 200 and "feedback_id" in r.text, "root 可导出用户反馈 CSV")
        r = self.c.delete(f"/api/chats/{self.alice_chat_id}")
        check(r.status_code == 204, "root 可删除其他用户问答")
        r = self.c.get(f"/api/chats/{self.alice_chat_id}")
        check(r.status_code == 404, "root 删除后问答不可再读取")
        r = self.c.get("/api/admin/feedback")
        check(not any(f["chat_id"] == self.alice_chat_id for f in r.json()["items"]), "删除问答同时删除关联反馈")
        r = self.c.get("/api/admin/audit?action=chat_delete")
        check(r.status_code == 200 and r.json()["total"] >= 2, "删除操作保留 chat_delete 审计")
        check(self.c.get("/api/admin/departments").status_code == 404, "root 后台部门接口已移除")
        r = self.c.get("/api/admin/overview")
        check(r.status_code == 200 and r.json()["counts"]["users"] >= 2, "概览计数正常")
        r = self.c.patch(
            "/api/admin/settings",
            json={"top_k": 4, "queries_per_minute": 11, "max_concurrent_llm": 2},
        )
        check(r.status_code == 200 and r.json()["settings"]["top_k"] == 4, "root 可修改并持久化运行参数")

    # ---------- 6. 限流 ----------
    def test_rate_limit(self) -> None:
        print("\n== 限流 ==")
        r = self.c.post(
            "/api/admin/users",
            json={"username": "bob", "password": "bob123456", "role": "user"},
        )
        check(r.status_code == 201, "创建 bob")
        self.c.post("/api/login", json={"username": "bob", "password": "bob123456"})
        hit = False
        for i in range(12):
            r = self.c.post("/api/query", json={"question": f"第 {i} 次测试问题"})
            if r.status_code == 429:
                hit = True
                break
        check(hit, "第 11 次问答触发 429 限流")

    def test_concurrency(self) -> None:
        print("\n== 5 用户并发问答 ==")
        r = self.login("root", ROOT_PW)
        check(r.status_code == 200, "并发测试前重新登录 root")
        accounts = [(f"load{i}", f"loadpass{i}") for i in range(5)]
        for username, password in accounts:
            r = self.c.post(
                "/api/admin/users",
                json={"username": username, "password": password, "role": "user"},
            )
            check(r.status_code == 201, f"创建并发用户 {username}")

        def ask(account: tuple[str, str]) -> int:
            username, password = account
            with httpx.Client(base_url=BASE_URL, timeout=30.0) as client:
                login = client.post("/api/login", json={"username": username, "password": password})
                if login.status_code != 200:
                    return login.status_code
                return client.post("/api/query", json={"question": f"{username} 查询住宿上限"}).status_code

        with ThreadPoolExecutor(max_workers=5) as pool:
            statuses = list(pool.map(ask, accounts))
        check(statuses == [200] * 5, f"5 用户并发均成功：{statuses}")

    # ---------- 7. LLM 异常不泄露 ----------
    def test_llm_errors(self) -> None:
        print("\n== LLM 异常 ==")
        r = self.login("alice", "alice123")
        check(r.status_code == 200, "alice 重新登录（避免被限流用户影响）")
        for trigger, code in [
            ("[[mock:http500]]", "llm_upstream"),
            ("[[mock:http401]]", "llm_auth"),
            ("[[mock:http429]]", "llm_rate_limited"),
            ("[[mock:sleep:2]]", "llm_timeout"),
        ]:
            r = self.c.post("/api/query", json={"question": trigger})
            body = r.json().get("detail", {})
            check(
                r.status_code == 502 and body.get("code") == code,
                f"{trigger} -> 502 code={body.get('code')}",
            )
            check(
                "mock-key" not in (r.text or "") and "Bearer" not in (r.text or ""),
                "异常响应不泄露 API Key",
            )
        # 无余额
        r = self.c.post("/api/query", json={"question": "[[mock:http402]]"})
        body = r.json().get("detail", {})
        check(r.status_code == 502 and body.get("code") == "llm_quota", "http402 -> llm_quota")

    # ---------- 8. 用户管理 ----------
    def test_user_admin(self) -> None:
        print("\n== 用户管理 ==")
        r = self.login("root", ROOT_PW)
        check(r.status_code == 200, "root 重新登录")
        users = self.c.get("/api/admin/users").json()["items"]
        alice = next(u for u in users if u["username"] == "alice")
        r = self.c.patch(f"/api/admin/users/{alice['id']}", json={"is_active": False})
        check(r.status_code == 200, "root 停用 alice")
        self.c.post("/api/logout")
        r = self.login("alice", "alice123")
        check(r.status_code == 403, "停用用户登录返回 403")
        r = self.login("root", ROOT_PW)
        check(r.status_code == 200, "root 再次登录")
        r = self.c.patch(f"/api/admin/users/{alice['id']}", json={"is_active": True})
        check(r.status_code == 200, "重新启用 alice")
        r = self.c.patch(
            f"/api/admin/users/{alice['id']}", json={"password": "newpass123"}
        )
        check(r.status_code == 200, "root 重置 alice 密码")
        with sqlite3.connect(os.environ["RAG_DB_PATH"]) as db:
            sessions = db.execute(
                "SELECT COUNT(*) FROM sessions WHERE user_id=?", (alice["id"],)
            ).fetchone()[0]
        check(sessions == 0, "重置密码后撤销 alice 的全部旧会话")

    def finish(self) -> None:
        r = self.login("root", ROOT_PW)
        r = self.c.post("/api/logout")
        self.c.close()
        print(f"\n结果: {len(PASS)} 通过, {len(FAIL)} 失败")
        if FAIL:
            print("失败项:")
            for f in FAIL:
                print(f"  - {f}")
            sys.exit(1)


def main() -> None:
    print(f"目标: {BASE_URL}")
    s = Smoke()
    s.test_auth()
    s.test_user_permissions()
    s.test_documents()
    s.test_kb_admin_permissions()
    s.test_query()
    s.test_history()
    s.test_admin_views()
    s.test_rate_limit()
    s.test_concurrency()
    s.test_llm_errors()
    s.test_user_admin()
    s.finish()


if __name__ == "__main__":
    main()
